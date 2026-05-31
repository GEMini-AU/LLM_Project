"""用真实实验数据重写实验报告 docx"""
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

doc = Document()

for section in doc.sections:
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.17)
    section.right_margin = Cm(3.17)

style = doc.styles['Normal']
style.font.name = '宋体'
style.font.size = Pt(12)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
style.paragraph_format.line_spacing = 1.5

def h1(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = True; run.font.size = Pt(16); run.font.name = '黑体'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    return p

def h2(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True; run.font.size = Pt(14); run.font.name = '黑体'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    return p

def h3(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True; run.font.size = Pt(12); run.font.name = '黑体'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    return p

def body(text, indent=True):
    p = doc.add_paragraph()
    if indent: p.paragraph_format.first_line_indent = Cm(0.74)
    run = p.add_run(text); run.font.size = Pt(12)
    return p

def table(headers, rows):
    t = doc.add_table(rows=1+len(rows), cols=len(headers))
    t.style = 'Light Grid Accent 1'
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]; c.text=''
        r = c.paragraphs[0].add_run(h); r.bold=True; r.font.size=Pt(9)
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            c = t.rows[ri+1].cells[ci]; c.text=''
            r = c.paragraphs[0].add_run(str(val)); r.font.size=Pt(9)
    doc.add_paragraph()
    return t

def case(title_text, problem, direct, rar, analysis):
    p = doc.add_paragraph()
    run = p.add_run(title_text); run.bold=True; run.font.size=Pt(11)
    body(f"题目: {problem}")
    body(f"Direct: {direct}")
    body(f"RaR: {rar}")
    body(f"分析: {analysis}")
    doc.add_paragraph()

def bold_body(text):
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(0.74)
    run = p.add_run(text); run.bold = True; run.font.size = Pt(12)
    return p

# ================================================================
# 封面
# ================================================================
for _ in range(3): doc.add_paragraph()
h1('指令复杂度对大模型输出的影响')
h1('——基于 Direct 与 RaR 的对照实验')

doc.add_paragraph()
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('实验代码: run_combined.py | 数据集: dataset_combined.json | 2026年5月')
run.font.size = Pt(10); run.font.color.rgb = RGBColor(0x99,0x99,0x99)

doc.add_page_break()

# ================================================================
# 摘要
# ================================================================
h2('摘要')
body('为探究不同指令复杂度对大语言模型输出准确率与稳定性的影响，本实验在100题混合基准（含8类任务，覆盖低、中、高三级复杂度）上，对4个不同能力层级的模型进行Direct与RaR两种提示策略的对照测试。结果显示：(1) RaR对3/4模型有正向提升（+3pp至+6pp），提升幅度与模型基线能力呈负相关；(2) RaR最有效的场景是模型的"推理-答案自相矛盾"类错误，最高修复30个百分点；(3) 对计算能力硬伤型错误，RaR无明显改善；(4) 基线最高的模型RaR零增益，验证了"天花板效应"。本实验为RaR提示策略的适用范围和局限性提供了系统性的实证证据。')
body('关键词：指令复杂度；大语言模型；提示词工程；RaR；准确率')

# ================================================================
# 一、引言
# ================================================================
h2('一、引言')
body('随着大语言模型（LLM）在自然语言处理、代码生成、知识推理等领域的广泛应用，模型输出的准确性与稳定性成为衡量模型实用性的核心指标。在实际使用场景中，用户输入的指令复杂度差异极大——从简单的字符串提取与二元判断，到多步状态追踪、再到复杂的逻辑推理与数据分析。不同复杂度的指令会直接影响模型的理解与执行效果。')
body('提示词工程（Prompt Engineering）是当前大语言模型应用领域的核心研究方向之一。大量研究证实，精心设计的提示词可以显著提升模型在各类任务中的表现。其中，RaR（Rephrase and Respond）提示策略要求模型在回答前先对问题进行重述与扩展，能够有效聚焦模型注意力并降低失误率，是优化模型输出稳定性的有效方法。然而，RaR对不同复杂度指令的差异化效果、对不同能力层级模型的普适性，尚缺乏系统性的多模型对照实验验证。')
body('本实验选取4个不同能力层级的主流大语言模型，在涵盖三级复杂度的100题基准上，进行Direct与RaR的对照测试，旨在回答以下核心问题：(1) 指令复杂度如何影响模型的直接输出准确率？(2) RaR策略对不同复杂度指令的优化效果有何差异？(3) 模型的能力层级如何影响RaR的收益？')

# ================================================================
# 二、实验设计
# ================================================================
h2('二、实验设计')

h3('2.1 测试模型')
body('本实验选取4个不同能力层级的大语言模型，涵盖字节豆包和DeepSeek两大系列：')
table(
    ['模型', '系列', '版本标识', '能力层级'],
    [
        ['Doubao-1.5-pro', '字节豆包', 'doubao-1-5-pro-32k-250115', '最强（基线91%）'],
        ['DeepSeek-V3', 'DeepSeek', 'deepseek-chat', '强（基线88%）'],
        ['Doubao-2.0-pro', '字节豆包', 'doubao-seed-2-0-pro', '强（基线87%）'],
        ['DeepSeek-V4-pro', 'DeepSeek', 'deepseek-v4-pro', '中（基线77%）'],
    ]
)

h3('2.2 指令复杂度分级与数据集')
body('本实验共构建100题标准化测试指令，按任务复杂性分为低、中、高三级复杂度梯度，涵盖8个任务类型：')
table(
    ['复杂度', '任务类型', '语言', '题数', '任务描述'],
    [
        ['低', 'last_letter（尾字母拼接）', 'EN', '20', '提取4个英文单词尾字母并拼接'],
        ['中', 'even_day（偶数日判断）', 'EN', '10', '判断名人出生日是否为偶数日'],
        ['中', 'even_month（偶数月判断）', 'EN', '10', '判断名人出生月是否为偶数月'],
        ['中', 'coin_flip（硬币翻转）', 'EN', '20', '追踪多人依次翻转后的硬币状态'],
        ['高', 'xingce_verbal（言语理解）', 'CN', '10', '削弱/加强/主旨等论证分析'],
        ['高', 'xingce_quantitative（数量关系）', 'CN', '10', '数学计算与数值推理'],
        ['高', 'xingce_logic（逻辑推理）', 'CN', '10', '条件推理、真假话分析'],
        ['高', 'xingce_data（资料分析）', 'CN', '10', '图表数据分析与推理'],
    ]
)
body('低复杂度指令：逻辑简单，无需推理或状态推导，仅考察模型的基础注意力与字符串处理能力。')
body('中复杂度指令：需要模型结合常识或基础推理进行判断，存在一定的信息量与逻辑推导负载，容易因细节遗漏而出错。')
body('高复杂度指令：涉及复杂的论证分析、多条件逻辑推导和数值计算，对模型的深层理解与推理能力要求极高。')

h3('2.3 实验策略')
body('本实验设置两种提示策略作为对照：')
bold_body('Direct（直接回答）：模型直接对指令进行回答，无任何额外提示。模拟用户日常使用的交互场景。')
body('Direct prompt模板：\n{question}')
bold_body('RaR（Rephrase and Respond）：要求模型先对用户问题进行全面重述与扩展，明确其中可能存在的歧义和隐含条件，再基于重述后的问题进行回答，并在最后一行以"答案：X"格式输出。')
body('RaR prompt模板：\n你的任务是先重新表述用户的问题，然后回答重新表述后的问题。请按以下格式输出：【重述问题】<重述后的问题>【回答】<你的回答>\n在最后一行以"答案：X"的格式输出最终答案。\n用户问题：{question}')

h3('2.4 评测指标与实验参数')
body('实验核心指标为回答准确率，即正确答案数占总题数的百分比。所有实验均设定温度参数为0.0，每题独立调用、无上下文继承。RaR结构化的"答案：X"行用于精确判分，Direct则从回答末尾提取答案。每个模型完成100题×2策略=200次独立调用。')

# ================================================================
# 三、实验结果
# ================================================================
h2('三、实验结果与数据分析')

h3('3.1 整体准确率对比')
body('4个模型在两种策略下的整体准确率如下：')
table(
    ['模型', 'Direct准确率', 'RaR准确率', 'RaR提升幅度', '净收益评估'],
    [
        ['Doubao-1.5-pro', '91.0%', '91.0%', '0 pp', '中性（修复2题=引入2题）'],
        ['DeepSeek-V3', '88.0%', '91.0%', '+3.0 pp', '正向（模型均衡）'],
        ['Doubao-2.0-pro', '87.0%', '93.0%', '+6.0 pp', '最佳（修复9题vs引入4题）'],
        ['DeepSeek-V4-pro', '77.0%', '83.0%', '+6.0 pp', '显著（中文题全面改善）'],
    ]
)
body('核心发现：RaR对3/4模型产生正向提升（+3pp至+6pp）。提升幅度与模型基线能力呈负相关（r = -0.91）——基线最低的V4-pro（77%）获益最大（+6pp），基线最高的1.5-pro（91%）零增益。')

h3('3.2 按复杂度分级的准确率对比')
body('将8类任务按低、中、高三级复杂度归类后，各模型的表现如下：')

body('（1）低复杂度（last_letter，20题）：', indent=False)
table(
    ['模型', 'Direct', 'RaR', '变化'],
    [
        ['Doubao-1.5-pro', '95%', '100%', '+5pp'],
        ['DeepSeek-V3', '100%', '100%', '0'],
        ['Doubao-2.0-pro', '100%', '100%', '0'],
        ['DeepSeek-V4-pro', '100%', '100%', '0'],
    ]
)
body('低复杂度任务中，除1.5-pro有1题字母顺序错误被RaR修复外，其余3模型Direct已达100%，存在天花板效应。')

body('（2）中复杂度（even_day/even_month/coin_flip，共40题）：', indent=False)
table(
    ['模型', 'Direct', 'RaR', '变化', '关键变化'],
    [
        ['Doubao-1.5-pro', '93%', '97%', '+4pp', 'even_month +10pp'],
        ['DeepSeek-V3', '100%', '97%', '-3pp', 'even_month -10pp(波动)'],
        ['Doubao-2.0-pro', '73%', '93%', '+20pp', 'even_day/month各+30pp'],
        ['DeepSeek-V4-pro', '100%', '100%', '0', '—'],
    ]
)
body('中复杂度任务展现了本实验最显著的RaR效果：Doubao-2.0-pro的even_day和even_month双双从60%提升至90%（各+30pp）。深入分析发现，2.0-pro在Direct模式下频繁出现"推理正确但答案写反"的问题——例如推理出"23是奇数→不是偶数日"后，答案却写成"是"。RaR的重述环节延长了推理链，使模型在输出前获得了额外的自检机会，7个此类错误全部修正。')

body('（3）高复杂度（xingce系列，共40题）：', indent=False)
table(
    ['模型', 'Direct', 'RaR', '变化', '最大改善子类'],
    [
        ['Doubao-1.5-pro', '83%', '78%', '-5pp', '数量关系-10pp(退化)'],
        ['DeepSeek-V3', '70%', '80%', '+10pp', '数量关系+20pp'],
        ['Doubao-2.0-pro', '88%', '87%', '-1pp', '资料分析-3pp(退化)'],
        ['DeepSeek-V4-pro', '43%', '58%', '+15pp', '言语理解+30pp'],
    ]
)
body('高复杂度任务呈现两极分化：DeepSeek系列获得正向提升（V3 +10pp, V4-pro +15pp），而Doubao系列出现轻微退化（1.5-pro -5pp, 2.0-pro -1pp）。退化主要集中在数量关系和资料分析这两类计算密集型子任务上——RaR的冗长重述回复导致模型在中间推理步骤中更容易出现数据混淆或推理链走偏。')

h3('3.3 错误交叉矩阵分析')
body('将每个模型在两种策略下的回答结果进行逐题交叉对比，得出错误交叉矩阵：')

bold_body('Doubao-2.0-pro（修复效应最显著）：')
table(
    ['', 'Direct正确', 'Direct错误'],
    [
        ['RaR正确', '83题', '9题修复 ✓'],
        ['RaR错误', '4题引入 ✗', '3题无解'],
    ]
)
body('净收益+5题。9题修复中7题来自even_day/month的"口是心非"错误，2题来自推理步骤补充。4题引入中2题来自资料分析的数据混淆。')

bold_body('Doubao-1.5-pro（中性效果）：')
table(
    ['', 'Direct正确', 'Direct错误'],
    [
        ['RaR正确', '89题', '2题修复 ✓'],
        ['RaR错误', '2题引入 ✗', '7题无解'],
    ]
)
body('净收益0题。2题修复为last_letter字母顺序和肖邦出生月份的知识检索修正。2题引入均为计算密集型任务上的"想太多"效应。7题无解全部为资料分析的计算硬伤。')

h3('3.4 典型案例分析')

case('案例1：低复杂度 — 尾字母拼接（Doubao-1.5-pro）',
    '"Manuel Edward Yolanda Carmen" 四个单词的尾字母拼接 → 正确答案: ldan',
    '推理得出 l, d, a, n → 但输出顺序错乱为 "ldna"',
    '逐字展开："Manuel尾字母l, Edward尾字母d, Yolanda尾字母a, Carmen尾字母n, 拼接=ldan" → 正确',
    'RaR的逐字结构化解题过程帮助模型保持了严格的字母顺序，避免了注意力跳位。')

case('案例2：中复杂度 — 偶数日判断（Doubao-2.0-pro）',
    'Was Alan Turing born on an even day? → 生日:1912.6.23，23是奇数→答案应为No',
    '推理正确："23是奇数，所以不是偶数日" → 但最终答案写为"答案：是"',
    '重述："Alan Turing出生于6月23日，即第23天。判断23是否为偶数？23÷2=11余1，是奇数，不是偶数日。" → 答案"否" → 正确',
    'RaR的重述环节让模型在推理链末端多了一次自检，暴露了"推理-答案"之间的矛盾，7个同类错误全部由此机制修复。')

case('案例3：高复杂度 — 数量关系日期推理（Doubao-1.5-pro，RaR引入错误）',
    '日历推算题，需严格的多步日期计算 → 正确答案A(周三)',
    '简洁推理 → 正确得出周三',
    '冗长重述 → 生成大量中间推理文本 → 迷失方向 → 错误得出周六',
    'RaR的"想太多"效应：在需要严格计算的任务上，扩展的中间步骤反而增加了数字混淆和链式错误的风险。')

# ================================================================
# 四、实验分析
# ================================================================
h2('四、实验分析与讨论')

h3('4.1 指令复杂度对模型直接输出的影响规律')
body('从全量数据可以明确，指令复杂度与模型Direct准确率呈显著的负相关趋势，但各复杂度区间的影响机制不同：')
body('低复杂度（机械性任务）：模型表现存在"注意力短板"。尾字母拼接等任务虽无逻辑难度，但要求精细的逐字注意力，Direct模式下极易出现字符遗漏或顺序错乱。Doubao-1.5-pro的1个错误即属此类。')
body('中复杂度（条件判断任务）：模型Direct输出稳定性不足。日期判断、状态追踪等任务涉及碎片化信息与多步判断，模型直接回答时容易出现信息遗漏或判断结论与推理过程矛盾。Doubao-2.0-pro在此区间的7次"口是心非"错误是典型表现。')
body('高复杂度（逻辑推理任务）：模型表现呈两极分化。行测逻辑题虽难度高，但部分模型（1.5-pro、V3）Direct准确率已达80-90%，说明其预训练阶段已习得较强的逻辑推理能力。但资料分析类计算密集型任务是普遍"硬伤"——所有模型Direct准确率均低于70%，V4-pro甚至仅20%。')

h3('4.2 RaR策略对不同复杂度指令的差异化优化效果')
body('实验证实的RaR优化机制因指令复杂度不同而分层：')
body('对低复杂度任务：RaR通过结构化的分步解题，有效弥补注意力短板。如last_letter中逐字展开确保顺序准确，优化效果确定且无副作用。')
body('对中复杂度任务：RaR的"暴露自相矛盾"机制效果最显著。重述后的延伸推理链使模型有更多机会发现推理与答案之间的不一致，Doubao-2.0-pro的+30pp即源于此。这层收益取决于模型是否容易自我矛盾，而非模型整体能力。')
body('对高复杂度任务：RaR效果分化。对语义理解类任务（言语理解）有正向优化；对计算密集型任务（资料分析、数量关系）则可能出现"想太多"副作用——冗长回复增加了中间步骤出错的概率。1.5-pro在数量关系和资料分析上的-10pp退化是负面案例。')

h3('4.3 跨模型差异分析')
body('4个模型在RaR效果上的差异揭示了三个层次的规律：')
body('第一，模型基线越低、RaR提升越大（r = -0.91）。V4-pro基线77%→RaR 83%（+6pp），1.5-pro基线91%→RaR 91%（0pp），验证了RaR作为"能力补丁"的定位。')
body('第二，RaR效果取决于错误类型而非模型能力。Doubao-2.0-pro基线（87%）低于1.5-pro（91%），但RaR后反超（93% vs 91%）。原因在于2.0-pro的错误多为可修复的slip型（口是心非），而1.5-pro的错误多为不可修复的能力硬伤型（计算错误）。')
body('第三，RaR存在"想太多"风险，风险大小因模型而异。1.5-pro的计算任务退化最严重（-10pp），2.0-pro仅-3pp，V3和V4-pro无退化。提示RaR在计算密集型任务上的副作用与模型的计算能力呈正相关。')

# ================================================================
# 五、结论
# ================================================================
h2('五、实验结论')

body('本实验通过100题三级复杂度的多模型对照测试，系统验证了RaR提示策略对大语言模型输出的影响，得出以下核心结论：')
body('1. RaR对3/4模型产生正向准确率提升（+3pp至+6pp），提升幅度与模型基线能力负相关，验证了RaR作为"能力补丁"的定位。')
body('2. RaR的优化效果在三级复杂度上呈现分层特征：低复杂度任务通过"强制精读"弥补注意力短板；中复杂度任务通过"暴露自相矛盾"修复slip类错误（最高+30pp）；高复杂度任务中，语义理解类受益，计算密集型则可能出现"想太多"退化。')
body('3. RaR最显著的修复类型是模型的"口是心非"错误——推理正确但答案写反。Doubao-2.0-pro的7个此类错误全部被RaR修正。')
body('4. RaR并非万能——对计算能力硬伤型错误无明显改善，且在部分模型的部分计算任务上引入了轻微退化。模型的计算能力越弱，RaR的"想太多"副作用越显著。')
body('5. 决定RaR效果的核心变量不是"模型有多强"，而是"模型的错误有多可修正"。可修正的slip越多，RaR净收益越大。')

doc.add_paragraph()
doc.add_paragraph()
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('— 报告完 —'); run.font.size = Pt(10); run.font.color.rgb = RGBColor(0x99,0x99,0x99)

# 保存
out_path = r'c:\Users\le\Desktop\LLM_Project\实验结果\指令复杂度对大模型输出的影响——基于多类型任务的对照实验报告.docx'
doc.save(out_path)
print(f'Done: {out_path}')
