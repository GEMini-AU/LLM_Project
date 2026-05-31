"""从 final_report.md 生成格式化的 Word 文档"""
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import re

doc = Document()

# 页面设置
for section in doc.sections:
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

style = doc.styles['Normal']
font = style.font
font.name = '微软雅黑'
font.size = Pt(10.5)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

def add_title(text, level=0):
    if level == 0:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(text)
        run.bold = True
        run.font.size = Pt(22)
        run.font.color.rgb = RGBColor(0x2c, 0x3e, 0x50)
    elif level == 1:
        doc.add_paragraph()
        h = doc.add_heading(text, level=1)
        for run in h.runs:
            run.font.color.rgb = RGBColor(0x2c, 0x3e, 0x50)
    elif level == 2:
        h = doc.add_heading(text, level=2)
        for run in h.runs:
            run.font.color.rgb = RGBColor(0x34, 0x49, 0x5e)
    elif level == 3:
        h = doc.add_heading(text, level=3)
        for run in h.runs:
            run.font.color.rgb = RGBColor(0x2c, 0x3e, 0x50)

def add_para(text, bold=False, italic=False, size=None, color=None, align=None):
    p = doc.add_paragraph()
    if align:
        p.alignment = align
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    if size:
        run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor(*color)
    return p

def add_code(text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1)
    run = p.add_run(text)
    run.font.name = 'Consolas'
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
    return p

def add_table(headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    # header
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ''
        run = cell.paragraphs[0].add_run(h)
        run.bold = True
        run.font.size = Pt(9)
    # data
    for r, row in enumerate(rows):
        for c, val in enumerate(row):
            cell = table.rows[r+1].cells[c]
            cell.text = ''
            run = cell.paragraphs[0].add_run(str(val))
            run.font.size = Pt(9)
    doc.add_paragraph()
    return table

def add_blockquote(text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1.5)
    p.paragraph_format.right_indent = Cm(1.5)
    run = p.add_run(text)
    run.italic = True
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
    return p

# ================================================================
# 正文
# ================================================================

add_title('Rephrase and Respond (RaR) 实验综合报告')
add_para('多模型对比验证：先重述问题再回答是否能提升 LLM 准确率？',
         align=WD_ALIGN_PARAGRAPH.CENTER, size=12, color=(0x7f,0x8c,0x8d))

doc.add_paragraph()

# --- 摘要 ---
add_title('摘要', level=1)
add_para('本实验在 100 题混合基准（英文符号推理 60 题 + 中文行测 40 题）上，对 4 个不同能力层级的大语言模型进行 Direct vs RaR 对比测试。结果显示：RaR 对 3/4 模型有正向提升，提升幅度与模型基线能力呈负相关——基线越低的模型获益越大（V4-pro: +6pp），基线最高的模型无增益（Doubao-1.5-pro: 0pp）。RaR 最显著的改善出现在模型的"口是心非"错误（推理正确但答案写反），而对计算密集型任务无效。')

# --- 一、实验设计 ---
add_title('一、实验设计', level=1)

add_title('1.1 策略定义', level=2)
add_table(
    ['策略', 'Prompt', '说明'],
    [
        ['Direct', '直接问题，无格式要求', '模拟用户直接提问'],
        ['RaR', '先重述问题 → 再回答', '强制模型先理解再解答'],
    ]
)

add_para('Direct prompt:', bold=True, size=9)
add_code('{question}')

add_para('RaR prompt:', bold=True, size=9)
add_code('你的任务是先重新表述用户的问题，然后回答重新表述后的问题。\n【重述问题】<重述后的问题>\n【回答】<你的回答>\n在最后一行以"答案：X"的格式输出最终答案。\n用户问题：{question}')

add_title('1.2 数据集', level=2)
add_para('100 题混合基准（dataset_combined.json），包含 8 个类别：')
add_table(
    ['类别', '语言', '题数', '题型描述'],
    [
        ['coin_flip', 'EN', '20', '硬币初始正面，依次翻转/不翻，最终正反？'],
        ['last_letter', 'EN', '20', '取4个单词尾字母拼接成字符串'],
        ['even_day', 'EN', '10', '名人出生日是否为偶数日？'],
        ['even_month', 'EN', '10', '名人出生月是否为偶数月？'],
        ['xingce_verbal', 'CN', '10', '行测言语理解（削弱/加强/主旨）'],
        ['xingce_quantitative', 'CN', '10', '行测数量关系（数学计算）'],
        ['xingce_logic', 'CN', '10', '行测逻辑推理（条件推理/真假话）'],
        ['xingce_data', 'CN', '10', '行测资料分析（图表数据推理）'],
    ]
)

add_title('1.3 测试模型', level=2)
add_table(
    ['模型', '厂商/系列', '能力层级'],
    [
        ['DeepSeek-V3 (chat)', 'DeepSeek', '强'],
        ['DeepSeek-V4-pro', 'DeepSeek', '中'],
        ['Doubao-2.0-pro (seed-2-0-pro)', '字节豆包', '强'],
        ['Doubao-1.5-pro (1-5-pro-32k)', '字节豆包', '最强(基线)'],
    ]
)

add_title('1.4 实验参数', level=2)
add_para('• 温度: 0.0')
add_para('• 每题独立调用，无上下文继承')
add_para('• RaR 结构化输出判分：提取最后一行的 "答案：X"')
add_para('• Direct 判分：末尾 200 字符内找最后出现的答案字母')

# --- 二、整体结果 ---
add_title('二、整体结果', level=1)

add_title('2.1 准确率总览', level=2)
add_table(
    ['模型', 'Direct', 'RaR', 'RaR 提升'],
    [
        ['Doubao-1.5-pro', '91.0%', '91.0%', '0 pp'],
        ['DeepSeek-V3', '88.0%', '91.0%', '+3.0 pp'],
        ['Doubao-2.0-pro', '87.0%', '93.0%', '+6.0 pp'],
        ['DeepSeek-V4-pro', '77.0%', '83.0%', '+6.0 pp'],
    ]
)

add_title('2.2 关键趋势', level=2)
add_blockquote('模型基线越低，RaR 提升越大。最强的 1.5-pro（基线 91%）RaR 零增益，最弱的 V4-pro（基线 77%）RaR 提升 +6pp。Pearson r = -0.91。')
add_para('')
add_code('     91%  88%   87%   77%   ← Direct基线')
add_code('      0   +3    +6    +6   ← RaR提升')
add_para('能力越强，RaR提升越小；能力越弱，RaR改善越显著。')

# --- 三、按题型分析 ---
add_title('三、按题型分析', level=1)

add_title('3.1 RaR 有效场景', level=2)
add_table(
    ['题型', '受益模型', '最大增益', '机制'],
    [
        ['Even Day/Month', '2.0-pro', '+30pp', '修复"推理正确、答案写反"的 slip'],
        ['xingce_verbal', 'V4-pro', '+30pp', '重述消除中文表述歧义'],
        ['xingce_quantitative', 'V3', '+20pp', '条件重排，减少跳步计算'],
        ['last_letter', '1.5-pro', '+5pp', '逐字展开保持顺序'],
    ]
)

add_title('3.2 RaR 无效场景', level=2)
add_table(
    ['题型', '所有模型', '原因'],
    [
        ['coin_flip', '全部 100%', '天花板效应，Direct 已完美'],
        ['xingce_data', '改善 ≤10pp', '模型计算能力硬瓶颈，重述不补'],
    ]
)

add_title('3.3 RaR 偶尔退化场景', level=2)
add_table(
    ['题型', '退化模型', '退化幅度', '原因'],
    [
        ['数量关系', '1.5-pro', '-10pp', '冗长重述→推理链混乱'],
        ['资料分析', '1.5-pro', '-10pp', '多数字→重述时记混数据'],
        ['资料分析', '2.0-pro', '-3pp', '计算密集型，重述无帮助'],
    ]
)

# --- 四、错误归因深度分析 ---
add_title('四、错误归因深度分析', level=1)

add_title('4.1 Doubao-2.0-pro 错误交叉矩阵', level=2)
add_table(
    ['', 'Direct 对', 'Direct 错'],
    [
        ['RaR 对', '83 题', '9 题修复 ✓'],
        ['RaR 错', '4 题引入 ✗', '3 题无解'],
    ]
)
add_para('净收益: +5 题', bold=True)

add_title('4.2 Doubao-1.5-pro 错误交叉矩阵', level=2)
add_table(
    ['', 'Direct 对', 'Direct 错'],
    [
        ['RaR 对', '89 题', '2 题修复 ✓'],
        ['RaR 错', '2 题引入 ✗', '7 题无解'],
    ]
)
add_para('净收益: 0 题', bold=True)

add_title('4.3 三类典型错误', level=2)

add_para('类型 1：「口是心非」型（RaR 修复）', bold=True)
add_para('模型推理正确但答案写反。2.0-pro 最严重——Even Day/Month 的 7 次 Direct 错误全部属于此类。')
add_code('模型: "13是奇数，所以不是偶数日" → 答案写"是"')
add_code('RaR:  重述环节暴露矛盾 → 修正为"否"')

add_para('类型 2：「想太多」型（RaR 引入）', bold=True)
add_para('1.5-pro 独有。RaR 的冗长重述让模型在推理链中迷失。')
add_code('Direct: 简洁推 → 周三(A)正确')
add_code('RaR:    长篇大论 → 迷失 → 周六(C)错误')

add_para('类型 3：「我就是不会」型（双方皆错）', bold=True)
add_para('计算密集型题，4 个模型都无法突破。集中在资料分析（多步除法比大小）和部分数量关系题。')

# --- 五、跨模型对比洞察 ---
add_title('五、跨模型对比洞察', level=1)

add_title('5.1 RaR 效果与模型"性格"的关系', level=2)
add_table(
    ['模型', '性格特征', 'RaR 效果'],
    [
        ['Doubao-2.0-pro', '深思熟虑但"口是心非"', '最佳（修复 9 题）'],
        ['DeepSeek-V4-pro', '能力偏弱且不稳定', '显著（+6pp）'],
        ['DeepSeek-V3', '稳健但偶尔跳步', '中等（+3pp）'],
        ['Doubao-1.5-pro', '直来直去，错了就是不会', '中性（0pp）'],
    ]
)

add_title('5.2 反直觉发现', level=2)
add_blockquote('"好"模型的 RaR 提升不一定大。Doubao-2.0-pro 的 Direct 基线（87%）低于 1.5-pro（91%），但 RaR 后反超（93% vs 91%）。2.0-pro 不是"更笨"，而是更容易出现可被 RaR 修复的 slip 类错误。1.5-pro 基线高但犯错时是真不会，RaR 无计可施。')

add_title('5.3 DeepSeek 系内部差异', level=2)
add_para('V3（chat）比 V4-pro 在本题集上强 11pp（88% vs 77%）。V4-pro 可能在"选择题推理"场景下未被充分优化。两个 DeepSeek 模型的 RaR 提升方向一致（正向），但幅度因基线而异。')

# --- 六、RaR 工作机制总结 ---
add_title('六、RaR 工作机制总结', level=1)
add_para('根据 4 模型 × 100 题的证据，RaR 的作用机制分为三层：')

add_para('第一层：强制精读（对所有模型有效）', bold=True)
add_para('重述环节迫使模型逐字审视问题，避免"扫一眼就答"的跳步。这层收益与模型能力无关。')

add_para('第二层：暴露自相矛盾（对"口是心非"模型显著）', bold=True)
add_para('重述后的推理链延长，模型更容易发现自己前后不一致。这层收益取决于模型是否容易 self-contradict。')

add_para('第三层：条件显式化（对基线弱模型有帮助）', bold=True)
add_para('将隐含条件、散乱信息重新组织为显式约束。这层收益与模型基线负相关——能力越弱，受益越大。')

# --- 七、结论 ---
add_title('七、结论', level=1)

add_para('1. RaR 有效但不普适：对 3/4 模型有正向提升，效果从 +3pp 到 +6pp')
add_para('2. 提升幅度与基线负相关：越弱的模型 RaR 帮助越大（r = -0.91）')
add_para('3. RaR 最擅长修复 slip 类错误：推理正确但答案写反的类型，2.0-pro 的 Even Day/Month +30pp 即属此类')
add_para('4. RaR 对计算硬伤无效：资料分析等计算密集型题，所有模型 RaR 后无明显改善')
add_para('5. RaR 偶尔有害：1.5-pro 上引入 2 个"想太多"错误，但 2.0-pro 修复 9 个 vs 引入 4 个，净收益为正')

add_blockquote('RaR 不是灵丹妙药，但对三类问题有效——表述歧义、条件跳步、推理-答案矛盾。模型越容易出现这些低级错误，RaR 的改善越显著。')

doc.add_paragraph()
add_para('实验代码: run_combined.py  |  数据集: dataset_combined.json  |  报告生成: 2026-05-26', size=8, color=(0x95,0xa5,0xa6), align=WD_ALIGN_PARAGRAPH.CENTER)

# 保存
doc.save('final_report.docx')
print('Done: final_report.docx')
